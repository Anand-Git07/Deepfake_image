from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/raviraj/Documents/New project/Data_Analytics_using_Microsoft_Tools_Internship_Report_Anushka_Sharma.docx")

STUDENT_NAME = "Anushka Sharma"
REGISTER_NO = "________________"
COLLEGE = "Administrative Management College"
COLLEGE_ADDR = "18th Km, Bannerghatta Road, Bangalore-560083"
DEPARTMENT = "Department of Computer Applications"
COURSE = "Bachelor of Computer Application"
SEMESTER = "VI Semester"
UNIVERSITY = "Bangalore University"
ACADEMIC_YEAR = "2025-2026"
TOPIC = "Data Analytics using Microsoft Tools"
COMPANY = "KSR Enterprises"
CERT_ISSUER = "S.L.A TRADERS"
INTERNSHIP_START = "2nd March 2026"
INTERNSHIP_END = "3rd April 2026"
CERT_DATE = "06-04-2026"
PLACE = "Bangalore"
GUIDE = "Mr. SANTHOSH KUMAR S"
GUIDE_TITLE = "Assistant Professor"
HOD = "Dr. Aruna Akula"
PRINCIPAL = "Dr. DVSSR PRAKASH"


def set_run_font(run, size=12, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def fix_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def paragraph(doc, text="", size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              before=0, after=6, line=1.5, color=None, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_together = keep
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def centered(doc, text, size=12, bold=False, after=6, before=0, color=None):
    return paragraph(doc, text, size=size, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER,
                     before=before, after=after, line=1.5, color=color)


def heading(doc, text, level=1, before=None):
    if level == 1:
        size, color, b, a = 16, "1F4D78", 10 if before is None else before, 8
    elif level == 2:
        size, color, b, a = 14, "1F4D78", 8 if before is None else before, 6
    else:
        size, color, b, a = 12, "365F91", 6 if before is None else before, 4
    return paragraph(doc, text, size=size, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                     before=b, after=a, line=1.5, color=color, keep=True)


def page_break(doc):
    doc.add_page_break()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = paragraph.add_run("Page ")
    set_run_font(r1, size=10)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.15)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    header = sec.header.paragraphs[0]
    header.text = f"{TOPIC} | {STUDENT_NAME}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, italic=True, color="666666")
    add_page_number(sec.footer.paragraphs[0])
    return doc


def add_cover(doc):
    centered(doc, "Internship Project Report on", size=16, bold=True, after=4, before=26)
    centered(doc, TOPIC, size=18, bold=True, after=18, color="1F4D78")
    centered(doc, "Submitted in partial fulfilment for the Award of the Internship", size=12, after=18)
    centered(doc, COURSE, size=14, bold=True, after=4)
    centered(doc, SEMESTER, size=13, bold=True, after=4)
    centered(doc, "2026", size=13, bold=True, after=18)
    centered(doc, UNIVERSITY, size=14, bold=True, after=28)
    centered(doc, "Under the Guidance", size=12, bold=True, after=8)
    centered(doc, "Internal Guide", size=12, bold=True, after=4)
    centered(doc, GUIDE, size=12, bold=True, after=2)
    centered(doc, GUIDE_TITLE, size=12, after=2)
    centered(doc, DEPARTMENT.upper(), size=12, bold=True, after=2)
    centered(doc, COLLEGE, size=12, bold=True, after=28)
    centered(doc, "SUBMITTED BY", size=12, bold=True, after=6)
    centered(doc, STUDENT_NAME, size=14, bold=True, after=2)
    centered(doc, f"Register No: {REGISTER_NO}", size=12, bold=True, after=0)
    page_break(doc)


def add_company_certificate(doc):
    centered(doc, "CERTIFICATE OF INTERNSHIP", size=16, bold=True, after=24, before=20)
    paragraph(doc, f"{CERT_DATE}", align=WD_ALIGN_PARAGRAPH.RIGHT, after=18)
    paragraph(doc, "TO WHOM IT MAY CONCERN", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    paragraphs = [
        f"This is to certify that Ms. {STUDENT_NAME}, a student of {COLLEGE}, Bannerghatta Road, studying {COURSE}, {SEMESTER}, affiliated to {UNIVERSITY}, Bangalore, has successfully completed a one-month Internship Programme on {TOPIC} at {COMPANY} from {INTERNSHIP_START} to {INTERNSHIP_END} at our office premises, Bengaluru.",
        f"During the tenure of her Internship Programme with us, Ms. {STUDENT_NAME} demonstrated a commendable work ethic and was found to be hardworking, sincere, and punctual. She actively contributed to assigned analytics exercises and exhibited a keen willingness to learn, adapt, and present business information through Microsoft-based reporting tools.",
        f"We wish Ms. {STUDENT_NAME} all the very best in her future academic and professional endeavours.",
    ]
    for t in paragraphs:
        paragraph(doc, t, after=10)
    paragraph(doc, "Yours faithfully,", after=16)
    paragraph(doc, f"For {COMPANY},", after=48)
    paragraph(doc, "Authorised Signatory", after=6)
    paragraph(doc, CERT_ISSUER, bold=True, after=10)
    paragraph(doc, f"This certificate is issued on behalf of {COMPANY} and is valid only with an authorised signature and company seal.", size=11, italic=True, after=0)
    page_break(doc)


def add_college_certificate(doc):
    centered(doc, COLLEGE.upper(), size=14, bold=True, after=2, before=12)
    centered(doc, COLLEGE_ADDR, size=12, after=2)
    centered(doc, DEPARTMENT.upper(), size=12, bold=True, after=14)
    centered(doc, COURSE.upper(), size=13, bold=True, after=12)
    centered(doc, "CERTIFICATE", size=16, bold=True, after=18)
    paragraph(doc, f'This is to certify that the Internship project titled "{TOPIC}" was successfully completed by {STUDENT_NAME.upper()} bearing register number {REGISTER_NO} of {SEMESTER} {COURSE} as prescribed by {UNIVERSITY} for the academic year {ACADEMIC_YEAR} under our guidance and supervision.', after=24)
    paragraph(doc, "Signature of the guide                                      Head of the Department", after=28)
    paragraph(doc, "(Signature with seal)", after=28)
    paragraph(doc, "Date of examination:", after=18)
    paragraph(doc, "Examiners:", after=18)
    paragraph(doc, "1.  _______________________", after=18)
    paragraph(doc, "2.  _______________________", after=0)
    page_break(doc)


def add_declaration(doc):
    centered(doc, "DECLARATION", size=16, bold=True, after=24, before=36)
    paragraph(doc, f'I hereby declare that the project being submitted entitled "{TOPIC.upper()}" for the partial fulfilment for the award of the "{COURSE}" degree is an authenticated record of work carried out by me in the 6th semester of BCA and submitted to {COLLEGE.upper()}, BANGALORE.', after=12)
    paragraph(doc, f"Under the guidance and supervision of {GUIDE}, {GUIDE_TITLE}, AMC.", after=42)
    paragraph(doc, "Thanking You,", after=42)
    paragraph(doc, STUDENT_NAME.upper(), bold=True, after=2)
    paragraph(doc, f"[{REGISTER_NO}]", after=0)
    page_break(doc)


def internship_day_rows():
    activities = [
        "Internship orientation, company introduction, and discussion of the data analytics internship objectives.",
        "Overview of business data sources, spreadsheet data types, and the role of clean data in reporting.",
        "Introduction to Microsoft Excel interface, workbook organisation, named ranges, and basic validation checks.",
        "Practice on data entry standards, duplicate identification, missing value checks, and column formatting.",
        "Learning sorting, filtering, conditional formatting, and quick exploratory analysis in Excel.",
        "Creating PivotTables to summarise monthly, category-wise, and region-wise business performance.",
        "Building PivotCharts and slicers for interactive exploration of summary tables.",
        "Introduction to Power Query for importing files, changing data types, and removing unwanted columns.",
        "Applying Power Query transformations: split columns, replace values, merge queries, and append data.",
        "Designing a simple data model with fact and dimension tables for reporting.",
        "Introduction to Power Pivot and relationships between tables using primary and foreign key fields.",
        "Learning DAX measures such as total value, average value, count of records, and percentage contribution.",
        "Creating time-based measures for month-to-date and year-to-date analysis using calendar tables.",
        "Introduction to Power BI Desktop workspace, report canvas, field pane, visual pane, and model view.",
        "Importing transformed Excel data into Power BI and validating loaded tables.",
        "Creating bar charts, line charts, cards, matrix visuals, and KPI indicators in Power BI.",
        "Designing slicers, filters, drill-down views, and cross-filtering interactions.",
        "Preparing an executive dashboard layout with summary cards, trend visuals, and category comparison.",
        "Applying colour rules, labels, titles, and formatting standards for readable business dashboards.",
        "Testing dashboard interactions and checking whether slicers update every visual correctly.",
        "Preparing insight notes from the dashboard and comparing Excel outputs with Power BI outputs.",
        "Learning report export, PDF sharing, and presentation of insights using Microsoft PowerPoint.",
        "Reviewing data privacy, file naming, documentation, and version control practices for analytics work.",
        "Drafting internship report chapters, literature survey notes, and methodology explanation.",
        "Preparing final report content, conclusion, and bibliography references.",
        "Final review of analytics workflow, internship learning outcomes, and report submission checklist.",
        "Discussion of improvements such as automated refresh, better data governance, and advanced DAX.",
        "Submission of final internship work summary to the guide and completion of internship activities.",
        "Completion of company-side formalities and collection of internship completion certificate.",
    ]
    d = date(2026, 3, 2)
    rows = []
    i = 0
    while d <= date(2026, 4, 3) and i < len(activities):
        if d.weekday() != 6:
            rows.append((str(i + 1), d.strftime("%d-%m-%Y"), activities[i]))
            i += 1
        d += timedelta(days=1)
    return rows


def add_day_book(doc):
    centered(doc, COLLEGE.upper(), size=13, bold=True, after=2)
    centered(doc, COLLEGE_ADDR, size=11, after=2)
    centered(doc, f"{DEPARTMENT} - BCA", size=12, bold=True, after=8)
    centered(doc, "Internship Day Book", size=16, bold=True, after=10)
    paragraph(doc, f"Student Name and Register No: {STUDENT_NAME.upper()} ({REGISTER_NO})", bold=True, after=2)
    paragraph(doc, f"Company Name: {COMPANY}", bold=True, after=2)
    paragraph(doc, "Address: Bengaluru", bold=True, after=8)
    rows = internship_day_rows()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].repeat_header = True
    headers = ["SL NO", "DATE", "ACTIVITIES / WORK DONE"]
    widths = [0.7, 1.25, 4.9]
    fix_table(table, widths)
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "D9EAF7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5)
    fix_table(table, widths)
    paragraph(doc, "", after=12)
    paragraph(doc, f"{GUIDE}                     {HOD}                                       {PRINCIPAL}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    paragraph(doc, "Assistant Professor                         HOD                                            Principal", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    paragraph(doc, "AMC COLLEGE                              AMC COLLEGE                                  AMC COLLEGE", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    page_break(doc)


def add_acknowledgement(doc):
    centered(doc, "ACKNOWLEDGEMENT", size=16, bold=True, after=24, before=20)
    texts = [
        "Apart from my efforts, the source of this internship project depends largely on the encouragement and guidance of many others. I take this opportunity to express my sincere gratitude to everyone who supported me during the successful completion of this internship report.",
        f"I would like to express my sincere thanks to our beloved Principal {PRINCIPAL}, who has been a leading light of our institution and has encouraged students to develop practical professional skills along with academic learning.",
        f"I would like to acknowledge my gratitude to our Head of the Department, {HOD}, {DEPARTMENT}, for her encouragement and support. Her guidance helped me understand the value of discipline, documentation, and continuous learning during the internship.",
        f"I would like to express my heartfelt gratitude to my guide {GUIDE}, {GUIDE_TITLE}, {DEPARTMENT}, for his valuable suggestions and support throughout this work. His guidance helped me shape the topic of {TOPIC} into a clear and useful academic report.",
        f"I also thank {COMPANY} for providing me the opportunity to complete an internship in the area of data analytics and Microsoft-based reporting. The exposure helped me understand how raw data can be converted into useful information for business decisions.",
        "Further, I would like to thank all professors of Computer Applications for their help and suggestions. I extend my special thanks to my parents and family members for their love and support.",
    ]
    for t in texts:
        paragraph(doc, t, after=8)
    paragraph(doc, f"Place: {PLACE}                                                                                                             {STUDENT_NAME.upper()}", after=2)
    paragraph(doc, f"({REGISTER_NO})", align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    page_break(doc)


def add_abstract(doc):
    centered(doc, "ABSTRACT", size=16, bold=True, after=18, before=18)
    texts = [
        f"The internship project titled {TOPIC} presents a structured study of how Microsoft Excel, Power Query, Power Pivot, Power BI, DAX, and Microsoft presentation tools can be used to convert raw organisational data into meaningful dashboards and reports. The project focuses on the complete analytics lifecycle: data collection, data cleaning, transformation, modelling, visualisation, interpretation, and reporting.",
        "The work demonstrates how Excel can be used for preliminary analysis through formulas, PivotTables, PivotCharts, and validation checks. Power Query is used to remove repeated manual cleaning steps by creating repeatable transformation logic. Power Pivot and DAX measures are used to define business calculations, while Power BI Desktop is used to design interactive dashboards that support faster decision-making.",
        "The project also highlights the importance of data quality, suitable chart selection, meaningful KPI design, dashboard layout, and responsible communication of insights. Rather than treating data analytics as only a technical exercise, the report explains how analytics connects business questions, data preparation, analytical modelling, and managerial reporting.",
        f"This internship at {COMPANY} helped me gain practical exposure to Microsoft-based analytics tools and understand how data can support everyday business monitoring. The report is prepared as a complete academic record of the internship, following the required format and front matter structure.",
    ]
    for t in texts:
        paragraph(doc, t, after=8)
    page_break(doc)


def add_contents(doc):
    centered(doc, "TABLE OF CONTENTS", size=16, bold=True, after=18, before=12)
    entries = [
        ("1.", "INTRODUCTION", "1-4"),
        ("2.", "COMPANY PROFILE", "5-6"),
        ("3.", "LITERATURE SURVEY", "7-11"),
        ("4.", "REQUIREMENTS SPECIFICATIONS", "12-16"),
        ("5.", "DATA ANALYTICS METHODOLOGY", "17-22"),
        ("6.", "MICROSOFT TOOL DESIGN", "23-28"),
        ("7.", "IMPLEMENTATION", "29-35"),
        ("8.", "ANALYSIS AND FINDINGS", "36-38"),
        ("9.", "SNAPSHOTS", "39-40"),
        ("10.", "INTERNSHIP LEARNING OUTCOMES", "41-42"),
        ("11.", "CONCLUSION", "43-44"),
        ("12.", "BIBLIOGRAPHY", "45"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    fix_table(table, [0.65, 5.2, 0.9])
    for i, h in enumerate(["No.", "Chapter", "Page No."]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=11, bold=True)
    for no, title, pages in entries:
        cells = table.add_row().cells
        for idx, value in enumerate([no, title, pages]):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(value)
            set_run_font(r, size=11, bold=(idx == 1))
    fix_table(table, [0.65, 5.2, 0.9])
    page_break(doc)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=12)


def add_chapter_page(doc, title, section_title, paragraphs, bullets=None, table=None, code=None):
    heading(doc, title, 1, before=0)
    heading(doc, section_title, 2, before=2)
    for t in paragraphs:
        paragraph(doc, t, after=6)
    if bullets:
        add_bullet_list(doc, bullets)
    if table:
        cols, rows, widths = table
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"
        for i, col in enumerate(cols):
            cell = tbl.rows[0].cells[i]
            set_cell_shading(cell, "E8F1FA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(col)
            set_run_font(r, size=10.5, bold=True)
        for row in rows:
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(val)
                set_run_font(r, size=9.5)
        fix_table(tbl, widths)
    if code:
        for line in code.splitlines():
            p = paragraph(doc, line, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.0)
            for run in p.runs:
                set_run_font(run, size=9.5, name="Courier New")
    page_break(doc)


def body_pages():
    return [
        ("INTRODUCTION", "1.1 Project Overview",
         ["Data Analytics using Microsoft Tools is an internship project that explains how data can be collected, prepared, analysed, visualised, and presented using tools from the Microsoft ecosystem. The project is designed around practical business reporting needs where raw data must be converted into a format that supports timely and accurate decisions.",
          "In many organisations, data is available in spreadsheets, transaction exports, attendance files, sales registers, customer lists, or operational logs. These files often contain duplicate records, inconsistent spellings, blank fields, incorrect data types, and multiple naming standards. The first responsibility of an analyst is therefore to prepare trustworthy data before creating any insight.",
          "Microsoft Excel, Power Query, Power Pivot, Power BI, and PowerPoint are useful because they allow an analyst to move from simple spreadsheet analysis to repeatable business intelligence reporting. Excel supports quick calculations and familiar tabular analysis, while Power BI supports interactive dashboards and visual storytelling."],
         ["To understand the end-to-end data analytics lifecycle using Microsoft tools.", "To create clean and repeatable data preparation steps using Power Query.", "To design simple dashboards that communicate KPIs, trends, and comparisons."]),
        ("INTRODUCTION", "1.2 Need for Data Analytics",
         ["The need for data analytics arises because business decisions are often made under time pressure. Managers require a clear view of what is happening, where performance is improving, and which areas require attention. A table of raw records rarely answers these questions directly, but structured analysis can convert the same records into meaningful indicators.",
          "Data analytics supports decisions by identifying patterns in historical data. For example, a monthly sales summary can show seasonal movement, while a regional comparison can identify locations that need additional focus. A dashboard can reduce the time spent manually searching through worksheets and allow users to concentrate on interpretation.",
          "The internship topic is therefore important for students of computer applications because it combines technical skill with business communication. The analyst must understand tools, data structures, formulas, and visual design, but must also present findings in a way that is easy for non-technical users to understand."],
         ["Analytics improves reporting speed and consistency.", "It reduces manual repetition by automating transformation steps.", "It helps users compare performance through dashboards and visual summaries."]),
        ("INTRODUCTION", "1.3 Objectives of the Project",
         ["The primary objective of this project is to understand how Microsoft tools can be used together for business data analytics. The project does not focus on one isolated tool; instead, it follows a workflow in which Excel supports initial analysis, Power Query supports data cleaning, Power Pivot and DAX support calculations, and Power BI supports dashboard creation.",
          "Another objective is to learn how to define meaningful KPIs. A report becomes useful only when the measures match the business question. Measures such as total revenue, average order value, growth percentage, completion rate, contribution percentage, and variance help users compare present performance with past performance or targets.",
          "The project also aims to develop professional reporting habits. These include keeping source files organised, documenting transformation steps, using clear column names, validating totals after every transformation, and designing dashboards that are readable, balanced, and focused."],
         ["Study the role of Microsoft Excel in data preparation and quick analysis.", "Use Power Query to clean, merge, and transform datasets.", "Create Power BI visuals and DAX measures for interactive reporting.", "Prepare an academic report explaining the complete workflow."]),
        ("INTRODUCTION", "1.4 Scope of the Project",
         ["The scope of the project covers the analytics lifecycle for a typical small or medium business dataset. It includes importing data, cleaning errors, transforming tables, preparing a data model, creating calculations, and designing dashboard pages. The report explains the process rather than claiming to build a production enterprise data warehouse.",
          "The project uses sample business scenarios such as sales performance, customer category analysis, region-wise reporting, product contribution, and monthly trend tracking. These scenarios are common in internships because they allow the analyst to practice real reporting concepts without depending on confidential company data.",
          "The scope also includes learning how to present findings. A dashboard is not complete only because visuals are placed on a canvas. It must have titles, filters, labels, suitable colours, and a logical reading order. The report therefore includes sections on visual design, interpretation, and future enhancements."],
         ["Data collection and preparation.", "Transformation using Power Query.", "KPI calculation using Excel formulas and DAX.", "Dashboard design using Power BI.", "Insight communication using Word and PowerPoint."]),
        ("COMPANY PROFILE", "2.1 Host Organisation",
         [f"The internship was completed at {COMPANY}, Bengaluru, as certified in the internship completion document. The company provided the practical environment for learning and applying the topic of {TOPIC}. The certificate confirms that the internship programme was conducted from {INTERNSHIP_START} to {INTERNSHIP_END}.",
          "The company environment helped in understanding how office data is maintained and used for reporting. Even when the datasets used for practice are simple, the important learning is that every record should be treated carefully because business decisions depend on the accuracy of the final report.",
          "During the internship, the focus was on developing a disciplined analytics workflow. The activities included spreadsheet preparation, Power Query transformation, PivotTable summaries, dashboard design, and report documentation."],
         ["Company Name: KSR Enterprises.", "Location: Bengaluru.", "Internship Area: Data Analytics using Microsoft Tools.", "Duration: 2nd March 2026 to 3rd April 2026."]),
        ("COMPANY PROFILE", "2.2 Relevance of Analytics in Business",
         ["Data analytics is useful for business organisations because it converts day-to-day transactions into managerial information. A company may maintain records of purchases, sales, stock, expenses, customers, or service requests, but these records become more valuable when they are summarised and interpreted.",
          "Microsoft tools are suitable for such environments because they are widely used in offices and can be adopted without heavy infrastructure. Excel is familiar to many employees, Power Query makes cleaning repeatable, and Power BI provides interactive dashboards that can be refreshed when new data is available.",
          "For an intern, this environment demonstrates the difference between academic exercises and workplace reporting. In academic work, sample data is usually clean. In real work, data quality must be checked and assumptions must be documented before presenting any number as final."],
         ["Data supports operational monitoring.", "Dashboards help compare branches, categories, and time periods.", "Repeatable transformations reduce manual errors.", "Visual reports help management understand trends quickly."]),
        ("COMPANY PROFILE", "2.3 Internship Role and Responsibilities",
         ["The internship role involved learning the tools and applying them to practical reporting situations. The work began with understanding the structure of data files and continued with cleaning, transformation, summarisation, and dashboard preparation. The role required attention to detail because a small formatting or formula error can affect the final result.",
          "A key responsibility was to maintain a logical workflow. This means keeping raw data separate from cleaned data, naming columns consistently, recording transformation steps, and validating totals after calculations. These habits are important because reports should be understandable even after some time has passed.",
          "Another responsibility was to communicate insights clearly. The final dashboard and report should not simply display many charts. They should answer what happened, why it matters, and which area deserves attention. This communication responsibility is a major part of data analytics work."],
         ["Clean and prepare spreadsheet data.", "Create PivotTables, PivotCharts, and Power BI visuals.", "Define basic measures and KPIs.", "Prepare documentation and internship report chapters."]),
        ("LITERATURE SURVEY", "3.1 Overview of Data Analytics",
         ["Data analytics is the process of examining data to identify patterns, relationships, trends, and useful conclusions. It is commonly divided into descriptive, diagnostic, predictive, and prescriptive analytics. This internship project mainly focuses on descriptive and diagnostic analytics because these are the foundation of reporting and dashboard development.",
          "Descriptive analytics answers the question of what happened. It summarises data using totals, averages, counts, percentages, and comparisons. Diagnostic analytics goes one step further by trying to understand why something happened. For example, if revenue decreased in a month, the analyst may check product category, region, customer group, and discount pattern.",
          "Microsoft tools support these analytics types through spreadsheet calculations, data models, and interactive visuals. Excel is often the starting point because it is accessible and flexible. Power BI extends the analysis by allowing multiple tables, relationships, slicers, and rich visual dashboards."],
         ["Descriptive analytics explains past performance.", "Diagnostic analytics investigates causes and differences.", "Dashboards help users explore the same data from multiple angles."]),
        ("LITERATURE SURVEY", "3.2 Microsoft Excel in Analytics",
         ["Microsoft Excel remains one of the most widely used tools for data analysis. Its importance comes from its flexibility, familiar interface, and support for formulas, tables, charts, and PivotTables. In small organisations, Excel is often the primary reporting tool because employees already understand workbook-based workflows.",
          "Excel helps analysts perform quick checks before building advanced dashboards. Functions such as SUMIFS, COUNTIFS, XLOOKUP, IFERROR, TEXT, DATE, and UNIQUE are useful for cleaning and summarising data. PivotTables allow users to group records by category, month, region, or other dimensions without writing complex code.",
          "However, Excel also requires discipline. If users manually edit formulas or copy files without version control, reports can become inconsistent. Therefore, good analytics practice uses structured tables, clear formulas, protected raw data, and documented assumptions."],
         None),
        ("LITERATURE SURVEY", "3.3 Power Query",
         ["Power Query is a data connection and transformation tool available in Excel and Power BI. It allows users to import data from different sources and apply cleaning steps through a repeatable query. This is important because cleaning data manually every month wastes time and increases the chance of error.",
          "Typical Power Query operations include removing blank rows, changing data types, renaming columns, splitting text fields, merging lookup tables, appending monthly files, filtering invalid records, and creating calculated columns. The applied steps pane records each transformation so that the analyst can review and modify the process.",
          "Power Query is especially useful when reports must be refreshed regularly. Once the transformation logic is prepared, new data can be loaded and the same steps can be applied again. This makes reporting more reliable than manual spreadsheet editing."],
         ["Import data from workbooks, CSV files, and folders.", "Clean data through recorded transformation steps.", "Combine multiple tables using merge and append operations.", "Refresh reports when the source file is updated."]),
        ("LITERATURE SURVEY", "3.4 Power BI and Business Intelligence",
         ["Power BI is a business intelligence tool used to create interactive dashboards and reports. It allows users to connect data, build relationships, define measures, and create visuals such as cards, bar charts, line charts, matrices, maps, and slicers. It is valuable because it helps business users interact with data instead of reading static tables.",
          "In Power BI, a data model can contain fact tables and dimension tables. The fact table contains measurable events such as sales or transactions, while dimensions contain descriptive fields such as customer, product, date, and region. This model helps users filter and aggregate data correctly.",
          "Power BI also supports DAX, which is a formula language used to create measures. Measures are calculated according to report filters and slicers, making them more flexible than fixed spreadsheet formulas. This allows the same measure to show values for different months, regions, or categories."],
         None),
        ("LITERATURE SURVEY", "3.5 Dashboard Design Principles",
         ["Dashboard design is not only about placing visuals on a page. A good dashboard has a clear purpose, a logical flow, and a visual hierarchy. Important KPIs should be visible first, supporting charts should explain movement or comparison, and filters should be easy to find without distracting from the main message.",
          "A dashboard should avoid unnecessary decoration and overcrowding. Too many colours, labels, or chart types can confuse the reader. Professional dashboards use consistent colours, meaningful titles, aligned visuals, and enough spacing. The chart type should match the data: line charts for trends, bar charts for comparisons, cards for single KPIs, and matrices for detailed lookup.",
          "The internship project applies these principles while designing Power BI report pages. Each page should answer a specific question such as overall performance, category contribution, region comparison, or monthly trend."],
         ["Keep the most important KPI at the top.", "Use slicers for date, category, and region filtering.", "Avoid chart clutter and repeated information.", "Use titles that explain what the visual shows."]),
        ("LITERATURE SURVEY", "3.6 Data Quality and Governance",
         ["Data quality is one of the most important factors in analytics. If source data contains errors, missing values, duplicates, or inconsistent categories, the final dashboard may show misleading results. Therefore, every analytics workflow should include validation before presenting findings.",
          "Common data quality checks include verifying required fields, checking date formats, comparing row counts before and after transformation, confirming that totals match between source and output, and reviewing unusual outliers. These checks build confidence in the report.",
          "Governance means maintaining rules for how data is named, stored, refreshed, and shared. Even in a student internship project, basic governance is useful. File names should be clear, raw data should not be overwritten, and assumptions should be documented in the report."],
         None),
        ("REQUIREMENTS SPECIFICATIONS", "4.1 Functional Requirements",
         ["The functional requirements define what the analytics solution must do. The project should allow the analyst to import data, clean the dataset, prepare summaries, create calculations, visualise important indicators, and present findings. The workflow must be understandable to users who are familiar with Microsoft Office tools.",
          "The system should support basic reporting questions such as total value, monthly performance, top categories, region-wise contribution, and performance against targets. It should also allow filtering so that users can explore a particular time period, category, or region.",
          "The final report and dashboard should present the business story clearly. Users should be able to understand what the data indicates without reading long technical explanations."],
         ["Import source data from Excel or CSV files.", "Clean and transform data using Power Query.", "Create PivotTables and DAX measures for KPIs.", "Build dashboard pages in Power BI.", "Export findings into report and presentation format."]),
        ("REQUIREMENTS SPECIFICATIONS", "4.2 Non-Functional Requirements",
         ["Non-functional requirements describe the quality of the solution. For this project, the analytics workflow should be accurate, maintainable, readable, and easy to refresh. Accuracy is essential because users may make decisions based on the final report.",
          "Maintainability means that another user should be able to understand the workbook, queries, measures, and dashboard layout. This requires meaningful names, organised folders, and simple documentation. Readability means that visuals and tables should not be crowded or confusing.",
          "The workflow should also be efficient. Power Query refreshes should complete in reasonable time for the size of the dataset, and dashboard pages should not contain unnecessary visuals that slow interaction."],
         ["Accuracy: calculated measures should match validated source totals.", "Usability: dashboard navigation should be simple and clear.", "Maintainability: transformations and formulas should be documented.", "Security: confidential data should be handled responsibly."]),
        ("REQUIREMENTS SPECIFICATIONS", "4.3 Hardware Requirements",
         ["The hardware requirements for this internship project are modest because the work is based on Microsoft desktop tools and sample business datasets. A normal student or office laptop can run Excel and Power BI Desktop if it has sufficient memory and storage.",
          "For larger datasets, additional RAM improves refresh speed and dashboard responsiveness. A stable internet connection is useful for documentation, template access, cloud storage, and optional publishing to Power BI Service, though the core project can be developed locally.",
          "The following table summarises the recommended hardware configuration for completing the project comfortably."],
         None,
         (["Component", "Recommended Specification"], [
             ["Processor", "Intel Core i3 / AMD equivalent or higher"],
             ["RAM", "8 GB recommended for Power BI Desktop"],
             ["Storage", "20 GB free disk space for data files and reports"],
             ["Display", "1366 x 768 minimum, Full HD recommended"],
             ["Network", "Broadband connection for updates and cloud sharing"],
         ], [2.0, 4.85])),
        ("REQUIREMENTS SPECIFICATIONS", "4.4 Software Requirements",
         ["The software requirements include Microsoft applications used for spreadsheet analysis, data transformation, dashboard creation, and documentation. The project can be completed with Microsoft Excel and Power BI Desktop as the main tools, while Word and PowerPoint support final reporting and presentation.",
          "Power Query is available inside Excel and Power BI, making it possible to use the same transformation logic in both environments. Power Pivot is useful for managing relationships and measures inside Excel, while Power BI provides a richer reporting canvas.",
          "The following software configuration was considered for the project."],
         None,
         (["Software", "Purpose"], [
             ["Microsoft Excel", "Data entry, validation, formulas, PivotTables, and PivotCharts"],
             ["Power Query", "Data cleaning, shaping, merging, appending, and refresh"],
             ["Power Pivot", "Data model relationships and measures inside Excel"],
             ["Power BI Desktop", "Dashboard design, DAX measures, and interactive reporting"],
             ["Microsoft Word", "Internship report preparation and formatting"],
             ["Microsoft PowerPoint", "Insight presentation and summary slides"],
         ], [2.0, 4.85])),
        ("REQUIREMENTS SPECIFICATIONS", "4.5 Dataset Requirements",
         ["A data analytics project requires data that is structured enough for analysis but realistic enough to demonstrate cleaning problems. The sample dataset for this report may include fields such as date, region, product category, customer type, quantity, sales value, cost, discount, and payment mode.",
          "The dataset should contain both numerical and categorical fields. Numerical fields allow calculations such as totals, averages, growth, and margins. Categorical fields allow grouping by product, location, customer segment, or transaction type. Date fields allow monthly and yearly trend analysis.",
          "The dataset should also be checked for quality issues. Missing dates, blank category fields, negative values, repeated rows, and inconsistent spelling can affect results. These issues are addressed during the Power Query stage."],
         ["Minimum 500 transaction records for meaningful summaries.", "Date, category, region, amount, quantity, and customer fields.", "A lookup table for region or category mapping where required.", "A target table for comparing actual performance with planned values."]),
        ("REQUIREMENTS SPECIFICATIONS", "4.6 Data Security Requirements",
         ["Although the project is academic, data security must still be considered. Business records may include customer names, phone numbers, addresses, invoice numbers, or other sensitive information. Such fields should not be shown unnecessarily in dashboards or screenshots.",
          "The principle of minimum exposure should be followed. Only fields required for analysis should be imported into the report model. If personal details are not needed for KPIs, they can be removed or masked during transformation. This keeps the final dashboard focused and reduces privacy risk.",
          "The report also recommends keeping raw files separate from edited files. A protected folder structure helps avoid accidental changes to source data and supports better auditability."],
         None),
        ("DATA ANALYTICS METHODOLOGY", "5.1 Analytics Lifecycle",
         ["The analytics lifecycle followed in this project begins with understanding the business question. Before opening a tool, the analyst should know what the report must explain. Typical questions include which month performed best, which category contributed most, whether sales met the target, and where improvement is required.",
          "After the question is defined, the analyst collects the required data and checks whether the fields are suitable. The next stage is cleaning and transformation, where Power Query is used to correct data types, remove errors, standardise values, and combine tables.",
          "The cleaned data is then modelled, calculated, visualised, and interpreted. The final stage is communication, where insights are written in a clear report or presented through a dashboard."],
         ["Understand the business question.", "Collect and validate source data.", "Clean and transform data.", "Model relationships and create measures.", "Visualise, interpret, and present insights."]),
        ("DATA ANALYTICS METHODOLOGY", "5.2 Data Collection",
         ["Data collection is the foundation of the project. The analyst gathers data from available files such as Excel workbooks, CSV exports, attendance sheets, sales registers, or manually maintained records. Each file is reviewed to understand its column names, data types, row count, and period covered.",
          "During collection, the analyst should identify whether all required fields are available. For example, a sales report requires date, item, quantity, amount, and category. If a field such as region or target value is missing, an additional lookup table may be required.",
          "A good practice is to preserve the original file and create a working copy for analysis. This ensures that the raw record remains available for verification if any doubt arises later."],
         None),
        ("DATA ANALYTICS METHODOLOGY", "5.3 Data Cleaning",
         ["Data cleaning removes problems that can distort analysis. In Excel, cleaning may begin with checking blank rows, duplicated records, spelling inconsistencies, invalid dates, and incorrect number formats. Conditional formatting can highlight unusual values, while filters can quickly show blanks or errors.",
          "In Power Query, cleaning becomes repeatable. The analyst can remove empty rows, promote headers, trim extra spaces, change data types, replace incorrect values, and split combined fields. These steps are recorded and can be refreshed when new data arrives.",
          "Validation is essential after cleaning. Row counts, total amounts, and category totals should be compared with the source file to ensure that the transformation has not removed valid data."],
         ["Remove blank rows and irrelevant columns.", "Change data types for date and numeric fields.", "Standardise category names.", "Remove duplicates only after confirming business rules.", "Validate cleaned totals against source totals."]),
        ("DATA ANALYTICS METHODOLOGY", "5.4 Data Transformation",
         ["Transformation changes data into a structure suitable for reporting. Many source files are designed for human reading, not analytics. For example, monthly columns may need to be unpivoted into a single date column, or separate files may need to be appended into one transaction table.",
          "Power Query supports transformation through a step-by-step interface. The analyst can merge a sales table with a product master table, append monthly files from a folder, create calculated columns, or extract month and year from a date field.",
          "Transformation should produce tidy tables in which each row represents one record, each column represents one attribute, and each table has a clear purpose. This structure makes modelling and visualisation easier."],
         None),
        ("DATA ANALYTICS METHODOLOGY", "5.5 Data Modelling",
         ["Data modelling organises tables so that reports calculate correctly. In a simple model, one fact table stores transactions and several dimension tables describe products, customers, dates, or regions. Relationships connect these tables through key fields.",
          "A date table is especially important for time-based reporting. It allows the analyst to group records by month, quarter, and year, and to create measures such as month-to-date, year-to-date, previous month, and growth percentage.",
          "The star schema is recommended for Power BI because it keeps the model understandable. A central fact table connects to smaller dimension tables, reducing confusion and improving calculation behaviour."],
         ["Fact table: transaction-level business records.", "Dimension tables: product, customer, region, and date information.", "Relationships: one-to-many links from dimensions to fact table.", "Measures: DAX calculations evaluated according to filters."]),
        ("DATA ANALYTICS METHODOLOGY", "5.6 KPI Selection",
         ["A key performance indicator is a measure that summarises progress toward a business objective. In this project, KPIs may include total sales, total orders, average order value, gross margin, target achievement percentage, growth rate, and top category contribution.",
          "A KPI should be selected because it helps answer a business question. Showing too many KPIs can reduce clarity. The dashboard should highlight only the measures that are most relevant to the reader.",
          "Each KPI should have a clear formula and interpretation. For example, target achievement percentage should be calculated as actual value divided by target value. The report should explain whether a higher value is good, what period is being measured, and which filters affect the KPI."],
         None),
        ("DATA ANALYTICS METHODOLOGY", "5.7 Validation and Testing",
         ["Validation ensures that the dashboard is trustworthy. After loading data into Power BI, totals should be compared with Excel PivotTables or source summaries. If the same total appears differently in two tools, the analyst must identify the reason before presenting the report.",
          "Testing also includes checking slicers, filters, relationships, and measures. A region slicer should filter all relevant visuals, and a date slicer should not remove unrelated reference values incorrectly. Measures should respond properly to filters.",
          "The project uses validation as a continuous activity rather than a final step. Every major transformation or calculation is checked so that errors are caught early."],
         ["Compare Power BI totals with Excel totals.", "Test slicer interactions.", "Check blank and unknown categories.", "Review DAX measures under different filters.", "Confirm that visual titles match the selected metric."]),
        ("MICROSOFT TOOL DESIGN", "6.1 Excel Workbook Design",
         ["The Excel workbook is designed with separate sheets for raw data, cleaned data, lookup tables, calculations, and dashboard summaries. Separating these sheets prevents accidental changes and makes the workbook easier to understand.",
          "Structured Excel tables are used because they expand automatically when new rows are added. They also make formulas easier to read because column names can be referenced instead of cell ranges. This is useful for functions such as SUMIFS, COUNTIFS, and XLOOKUP.",
          "The workbook should use clear sheet names, consistent date formats, and protected formulas where required. Colours may be used lightly to distinguish input cells, calculated cells, and output summaries."],
         None),
        ("MICROSOFT TOOL DESIGN", "6.2 Excel Formula Design",
         ["Formulas are used for quick calculations and validation checks. SUMIFS can calculate totals for a selected category, COUNTIFS can count records matching conditions, XLOOKUP can retrieve values from a master table, and IFERROR can provide readable results when a lookup fails.",
          "Formula design should focus on clarity. A long formula that no one can understand is difficult to maintain. Where possible, helper columns or Power Query transformations can simplify calculations.",
          "The project uses formulas for validation and small calculations, while larger reusable measures are handled in Power BI using DAX."],
         None,
         None,
         """=SUMIFS(Sales[Amount], Sales[Region], G2, Sales[Month], H2)
=COUNTIFS(Sales[Category], "Electronics", Sales[Status], "Completed")
=XLOOKUP([@ProductCode], ProductMaster[Code], ProductMaster[Category], "Unknown")
=IFERROR([@Sales]/[@Target], 0)"""),
        ("MICROSOFT TOOL DESIGN", "6.3 Power Query Design",
         ["Power Query is designed as the main cleaning layer. The query starts by importing the source file, promoting headers, changing data types, and removing empty rows. It then applies transformations such as trimming spaces, replacing values, and merging lookup tables.",
          "A separate query is maintained for each source table. The final output query is named clearly so that the report user can identify which table is loaded into the model. Intermediate queries can be disabled from loading if they are used only for preparation.",
          "The main advantage of this design is repeatability. When the source file is updated, the analyst refreshes the query instead of performing the same cleaning steps manually."],
         None,
         None,
         """let
    Source = Excel.Workbook(File.Contents("sales_data.xlsx"), null, true),
    Sales_Table = Source{[Item="Sales",Kind="Table"]}[Data],
    ChangedType = Table.TransformColumnTypes(Sales_Table,{{"Date", type date}, {"Amount", type number}}),
    RemovedBlankRows = Table.SelectRows(ChangedType, each [Date] <> null),
    TrimmedText = Table.TransformColumns(RemovedBlankRows,{{"Region", Text.Trim, type text}})
in
    TrimmedText"""),
        ("MICROSOFT TOOL DESIGN", "6.4 Power BI Report Layout",
         ["The Power BI report layout follows a simple reading order. The top area contains KPI cards, the middle area contains trend and comparison visuals, and the bottom area contains detailed tables or supporting charts. Slicers are placed where users can find them easily.",
          "The dashboard uses consistent colours to avoid confusion. For example, one colour may represent actual performance and another may represent target performance. Negative values can be highlighted carefully, but excessive red or green should be avoided.",
          "Each visual has a clear title and avoids unnecessary gridlines. Labels are shown only where they improve understanding. This design makes the dashboard suitable for repeated business review."],
         None),
        ("MICROSOFT TOOL DESIGN", "6.5 DAX Measure Design",
         ["DAX measures are used in Power BI to calculate values that respond to filters. Unlike calculated columns, measures are evaluated dynamically based on slicers and visual context. This makes them useful for dashboards where users explore different periods or categories.",
          "Common measures include total sales, total quantity, average order value, previous month sales, growth percentage, and target achievement. These measures should be named clearly and grouped logically in the model.",
          "DAX formulas should be tested under different filters to ensure they behave correctly. A measure may show the correct grand total but incorrect category values if relationships or filter context are not properly handled."],
         None,
         None,
         """Total Sales = SUM(Sales[Amount])
Total Orders = DISTINCTCOUNT(Sales[Order ID])
Average Order Value = DIVIDE([Total Sales], [Total Orders])
Sales LY = CALCULATE([Total Sales], SAMEPERIODLASTYEAR('Date'[Date]))
Growth % = DIVIDE([Total Sales] - [Sales LY], [Sales LY])"""),
        ("MICROSOFT TOOL DESIGN", "6.6 Visual Selection",
         ["The choice of visual affects how quickly users understand the report. KPI cards are suitable for single important numbers, line charts are suitable for trends, bar charts are suitable for comparisons, and matrices are suitable for detailed lookup.",
          "Pie charts and donut charts should be used carefully because they become difficult to read when there are many categories. In most business dashboards, a sorted bar chart gives clearer comparison than a pie chart.",
          "The project uses a mix of cards, line charts, clustered bar charts, tables, and slicers. This combination supports both overview and detail without making the dashboard crowded."],
         None,
         (["Question", "Recommended Visual"], [
             ["What is the total performance?", "Card or KPI visual"],
             ["How did values change over time?", "Line chart"],
             ["Which category is highest?", "Sorted bar chart"],
             ["How does actual compare with target?", "Column chart or KPI card"],
             ["What are the detailed records?", "Table or matrix"],
         ], [3.2, 3.65])),
        ("MICROSOFT TOOL DESIGN", "6.7 Report Sharing and Documentation",
         ["After the dashboard is prepared, the findings must be documented. Documentation includes the data source, cleaning steps, formulas, assumptions, refresh process, and interpretation of important visuals. This makes the report easier to review and maintain.",
          "PowerPoint can be used to present a short executive summary, while Word is used for the complete internship report. If Power BI Service is available, reports can be published and shared with controlled access.",
          "The report should avoid unsupported claims. Every insight should come from a measure, chart, or validated observation. This habit improves the professional quality of analytics communication."],
         None),
        ("IMPLEMENTATION", "7.1 Preparing Source Data",
         ["The implementation begins by preparing source data in Excel. The raw sheet is kept unchanged, and a working sheet or Power Query connection is used for analysis. Column names are reviewed and changed to meaningful names such as Date, Region, Product, Quantity, Sales, Cost, and Category.",
          "The analyst checks whether dates are stored as actual date values and whether numeric columns are stored as numbers. Incorrect data types are a common reason for wrong calculations. The row count and total sales amount are noted before cleaning.",
          "The prepared source data becomes the input for Power Query. This ensures that cleaning steps are applied in a controlled way and can be repeated when data changes."],
         None),
        ("IMPLEMENTATION", "7.2 Cleaning with Power Query",
         ["In Power Query, the source table is loaded and each transformation is applied in sequence. Blank rows are removed, column types are corrected, unnecessary columns are deleted, and text fields are trimmed to remove extra spaces.",
          "Inconsistent category names are standardised. For example, if the same region appears as Bengaluru, Bangalore, and BLR, a replacement or mapping table can convert them into one standard value. This improves grouping accuracy in PivotTables and Power BI visuals.",
          "After cleaning, the query output is loaded to the data model. The analyst validates row count and totals to ensure that the query did not remove valid records."],
         None),
        ("IMPLEMENTATION", "7.3 Creating PivotTables",
         ["PivotTables are created in Excel to produce quick summaries. They are useful for validating totals before building the Power BI dashboard. A PivotTable can show sales by month, sales by region, quantity by product category, or count of transactions by payment mode.",
          "Slicers can be added to PivotTables to filter the report by month, region, or category. This gives an early interactive view of the data and helps identify which dimensions are useful for the final dashboard.",
          "PivotTables also help detect unusual values. If a blank category appears as a separate row or if a region has an unexpectedly high value, the analyst can return to Power Query and fix the source issue."],
         None),
        ("IMPLEMENTATION", "7.4 Building the Data Model",
         ["The Power BI data model is created by importing the cleaned table and related lookup tables. Relationships are defined between the transaction table and dimension tables. A date table is added for calendar-based reporting.",
          "The model is checked in the relationship view to ensure that filters flow from dimension tables to the fact table. Many-to-many relationships are avoided unless necessary because they can create confusing results.",
          "A clean model improves dashboard reliability. When users filter by month, region, or category, all visuals should respond consistently because the relationships are properly defined."],
         None),
        ("IMPLEMENTATION", "7.5 Creating DAX Measures",
         ["DAX measures are created for the main KPIs. The first measure is total sales, followed by total orders, total quantity, average order value, target achievement, and growth percentage. These measures are added to visuals rather than using raw columns directly.",
          "Using measures improves consistency. If total sales is defined once, every visual uses the same calculation. This avoids the risk of one chart using a different formula from another chart.",
          "The measures are tested by placing them in a matrix with month and region fields. This confirms whether the values change correctly under filter context."],
         None),
        ("IMPLEMENTATION", "7.6 Designing Dashboard Page One",
         ["The first dashboard page is designed as an executive overview. It contains KPI cards for total sales, total orders, average order value, and target achievement. A line chart shows monthly trend, while a bar chart shows region-wise comparison.",
          "The purpose of this page is to answer the first-level question: how is overall performance? Users should be able to understand the current position within a few seconds. Detailed tables are avoided on this page to keep the overview clean.",
          "Slicers for month, region, and category are added so that users can focus on a specific view. The slicers are formatted consistently and placed in the same area."],
         None),
        ("IMPLEMENTATION", "7.7 Designing Dashboard Page Two",
         ["The second dashboard page focuses on category and product analysis. It shows contribution by product category, top products, low-performing products, and quantity movement. This page helps identify which categories drive the main results.",
          "A sorted bar chart is used for category contribution because it makes comparison easy. A table or matrix shows product-level detail for users who need more information. Conditional formatting can highlight values above or below target.",
          "This page demonstrates how analytics can move from overall performance to specific areas that need action."],
         None),
        ("IMPLEMENTATION", "7.8 Designing Dashboard Page Three",
         ["The third dashboard page focuses on customer or regional analysis depending on the available fields. It can show performance by region, customer type, or payment mode. The goal is to identify where performance differs and which segment requires attention.",
          "A map visual may be used if geographic data is reliable, but a bar chart is often clearer for comparing regions. The report avoids complicated visuals when simple charts communicate the message better.",
          "The page also includes a short insight section that summarises what the user should notice. This connects the visuals with business interpretation."],
         None),
        ("IMPLEMENTATION", "7.9 Dashboard Testing",
         ["Dashboard testing is performed after visuals are created. The analyst checks whether slicers filter all relevant visuals, whether KPI cards match matrix totals, and whether charts display correct date ordering. Any unexpected blank category is investigated.",
          "Testing also checks layout quality. Titles should not be cut, labels should be readable, and colours should be consistent. Visuals should not overlap, and the dashboard should remain useful when filters are applied.",
          "The final dashboard is reviewed by comparing selected values with Excel summaries. This cross-check gives confidence that the Power BI report is accurate."],
         None),
        ("IMPLEMENTATION", "7.10 Preparing Insight Notes",
         ["Insight notes translate dashboard observations into meaningful statements. For example, instead of saying that a bar is taller, the note should explain which region contributed most and whether the difference is significant. This helps users focus on action.",
          "Good insight writing avoids unsupported claims. Every statement should be traceable to a chart or calculation. If the data shows correlation but not cause, the report should not claim a definite reason without additional evidence.",
          "The internship project includes insight notes to demonstrate that analytics is not just tool operation. It is a process of interpreting data responsibly."],
         None),
        ("ANALYSIS AND FINDINGS", "8.1 Key Findings from the Dashboard",
         ["The dashboard shows how Microsoft tools can convert transactional data into clear performance summaries. Excel provides initial validation, Power Query improves data quality, and Power BI presents the results in a visual form. The same dataset becomes more useful when it is cleaned, modelled, and arranged around business questions.",
          "One key finding is that data quality issues appear before analysis begins. Inconsistent region names, blank fields, and wrong data types can create misleading totals. Power Query is effective in solving these issues because it records repeatable cleaning steps.",
          "Another finding is that dashboard layout affects interpretation. A clear dashboard with selected KPIs communicates faster than a crowded report with too many visuals."],
         None),
        ("ANALYSIS AND FINDINGS", "8.2 Performance Analysis",
         ["Performance analysis uses KPIs and comparisons to understand the dataset. Total sales shows overall volume, while average order value explains the typical transaction size. Region-wise and category-wise comparisons show where performance is concentrated.",
          "Monthly trend analysis helps identify whether performance is increasing, decreasing, or fluctuating. If a month shows a sudden drop, the analyst can filter by category or region to investigate the cause. This demonstrates the value of interactive reporting.",
          "The performance analysis also highlights the importance of target comparison. A high sales value may still be below target, while a smaller region may perform well against its planned objective."],
         None),
        ("ANALYSIS AND FINDINGS", "8.3 Data Quality Findings",
         ["During the cleaning stage, common data quality issues include blank entries, inconsistent text values, repeated records, and columns stored with incorrect data types. These issues may look small but can affect every chart that depends on the field.",
          "Power Query helps detect and correct many of these issues. The applied steps list creates transparency because each change can be reviewed. This is more reliable than manual editing across many worksheets.",
          "The project shows that data quality is not a separate activity but part of the analytics workflow. Every insight depends on the quality of the data used to create it."],
         ["Standardise spelling before grouping categories.", "Validate totals after removing duplicates.", "Use date tables to avoid wrong month sorting.", "Document assumptions for missing and corrected values."]),
        ("ANALYSIS AND FINDINGS", "8.4 Business Interpretation",
         ["Business interpretation is the stage where numbers are converted into meaning. A dashboard may show that one category is leading, but interpretation explains whether that is expected, whether it is growing, and whether it needs further attention.",
          "In this project, interpretation is supported by slicers and drill-down visuals. Users can move from total performance to month, region, and category views. This helps them ask follow-up questions without creating a new report each time.",
          "The interpretation section also teaches that analytics should be careful and honest. If the data does not include a reason for a trend, the analyst should present it as an observation and recommend further investigation."],
         None),
        ("SNAPSHOTS", "9.1 Excel and Power Query Screens",
         ["The following snapshot descriptions represent the main working screens used during the project. The Excel workbook contains raw data, cleaned data, PivotTable summaries, and validation checks. The Power Query editor contains applied steps such as changed type, removed blanks, trimmed text, and merged lookup table.",
          "In the final report, these screens demonstrate how the analysis moved from raw records to cleaned tables. They also show that transformation steps were not performed randomly but followed an organised process.",
          "Snapshot placeholders are included in the report structure so that actual screenshots can be inserted if required by the college or guide."],
         None,
         (["Snapshot", "Description"], [
             ["Excel Raw Data Sheet", "Transaction records before cleaning and validation"],
             ["Power Query Applied Steps", "Repeatable transformations used for data preparation"],
             ["PivotTable Summary", "Month-wise and region-wise validation summaries"],
         ], [2.2, 4.65])),
        ("SNAPSHOTS", "9.2 Power BI Dashboard Screens",
         ["The Power BI dashboard contains an executive overview page, a category analysis page, and a region analysis page. Each page uses slicers and visuals to answer a particular business question.",
          "The executive overview page includes KPI cards and trend charts. The category analysis page shows contribution and product comparison. The regional page shows performance across locations or segments depending on the available data.",
          "These dashboard screens complete the analytics workflow by presenting cleaned and calculated information in a form that can be used for decisions."],
         None,
         (["Dashboard Page", "Main Visuals"], [
             ["Executive Overview", "KPI cards, monthly trend, region comparison"],
             ["Category Analysis", "Category contribution, top product chart, matrix"],
             ["Regional Analysis", "Region-wise bar chart, slicers, insight notes"],
         ], [2.2, 4.65])),
        ("INTERNSHIP LEARNING OUTCOMES", "10.1 Technical Learning",
         ["The internship improved my understanding of Microsoft tools used for data analytics. I learned how Excel supports validation and quick analysis, how Power Query makes cleaning repeatable, and how Power BI converts data models into interactive dashboards.",
          "I also learned the importance of DAX measures and filter context. Measures such as total sales, average order value, and growth percentage become more powerful when they respond to slicers and visual filters.",
          "Another important learning outcome was the need for documentation. A dashboard should be accompanied by information about data sources, cleaning steps, formulas, assumptions, and limitations."],
         ["Excel formulas and PivotTables.", "Power Query transformations.", "Data modelling and relationships.", "DAX measures and Power BI visuals.", "Report writing and insight presentation."]),
        ("INTERNSHIP LEARNING OUTCOMES", "10.2 Professional Learning",
         ["The internship also developed professional qualities such as punctuality, attention to detail, communication, and willingness to learn. Data analytics work requires patience because errors are not always visible immediately. A good analyst checks results carefully before presenting them.",
          "I learned that technical output must be understandable to others. A dashboard is successful only when users can interpret it easily. Therefore, layout, titles, labels, and explanation are part of the professional responsibility of an analyst.",
          "The internship helped me connect classroom learning with practical reporting work. It strengthened my confidence in using Microsoft tools for solving business information problems."],
         None),
        ("CONCLUSION", "11.1 Conclusion",
         [f"The internship project on {TOPIC} successfully explains the complete process of converting raw data into useful business information using Excel, Power Query, Power Pivot, Power BI, DAX, Word, and PowerPoint. The project demonstrates how Microsoft tools can support data cleaning, modelling, calculation, visualisation, and reporting.",
          "The most important learning from the project is that analytics depends on both technical skill and careful thinking. A dashboard is meaningful only when the data is clean, the measures are correct, and the visual design helps users understand the message. The project also shows that repeatable transformation steps reduce manual errors and save time.",
          f"This internship at {COMPANY} provided practical exposure to data analytics activities and improved my ability to prepare professional reports. The knowledge gained through this work will be useful for future academic projects and professional opportunities in analytics and business intelligence."],
         None),
        ("CONCLUSION", "11.2 Future Enhancements",
         ["The current project can be extended by connecting Power BI directly to live databases or cloud storage instead of using static files. Automated refresh can reduce manual effort and keep dashboards updated.",
          "More advanced DAX measures can be added for forecasting, customer segmentation, moving averages, and variance analysis. Power BI Service can be used to publish dashboards securely and share them with selected users.",
          "The project can also include stronger data governance practices such as role-based access, data dictionaries, version history, and documented refresh schedules. These improvements would make the solution closer to professional business intelligence practice."],
         ["Automated refresh from cloud folders or databases.", "Advanced DAX and forecasting measures.", "Power BI Service publishing and access control.", "Data dictionary and refresh documentation.", "Integration with Microsoft Teams or SharePoint for collaboration."]),
        ("BIBLIOGRAPHY", "12.1 References",
         ["Microsoft Support. Microsoft Excel documentation and help resources.",
          "Microsoft Learn. Power Query documentation for Excel and Power BI.",
          "Microsoft Learn. Power BI Desktop documentation and data modelling guidance.",
          "Microsoft Learn. Data Analysis Expressions (DAX) reference.",
          "Microsoft Learn. Power BI visualisation and report design guidance.",
          "Few, Stephen. Information Dashboard Design: Displaying Data for At-a-Glance Monitoring. Analytics Press.",
          "Provost, Foster and Fawcett, Tom. Data Science for Business. O'Reilly Media.",
          "Kimball, Ralph and Ross, Margy. The Data Warehouse Toolkit. Wiley.",
          "Administrative Management College internship report format reference supplied by the student.",
          f"Internship completion certificate issued for {STUDENT_NAME} by {COMPANY} / {CERT_ISSUER}, dated {CERT_DATE}."],
         None),
    ]


def add_body(doc):
    omitted_to_keep_report_near_55_pages = {
        "2.2 Relevance of Analytics in Business",
        "3.6 Data Quality and Governance",
        "4.6 Data Security Requirements",
        "5.7 Validation and Testing",
        "6.7 Report Sharing and Documentation",
        "7.7 Designing Dashboard Page Two",
        "7.8 Designing Dashboard Page Three",
        "7.10 Preparing Insight Notes",
        "8.4 Business Interpretation",
    }
    for spec in body_pages():
        if spec[1] in omitted_to_keep_report_near_55_pages:
            continue
        add_chapter_page(doc, *spec)


def main():
    doc = setup_doc()
    add_cover(doc)
    add_company_certificate(doc)
    add_college_certificate(doc)
    add_declaration(doc)
    add_day_book(doc)
    add_acknowledgement(doc)
    add_abstract(doc)
    add_contents(doc)
    add_body(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
